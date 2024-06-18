import os
from tempfile import gettempdir
from pypdf import PdfReader
from pydub import AudioSegment
from openai import OpenAI
import json
from docx import Document
from docx2pdf import convert
import re
from dotenv import load_dotenv
from pathlib import Path

load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
MODEL = "gpt-4o"
client = OpenAI(api_key=OPENAI_API_KEY)

def create_chapters_for_title(title):
    response = client.chat.completions.create(
        model=MODEL,
        response_format={"type": "json_object"},
        max_tokens=3000,
        messages=[
            {"role": "system", "content": "You are a creative eBook writer. Give the output as JSON."},
            {"role": "user", "content": f"Write the chapters and subheadings for the ebook titled '{title}'. Make the chapter names as keys and the list of subheadings as values. Give 4 chapters and 3 subheadings for each chapter."} #Change the number of chapters and subchapter as desired
        ]
    )
    with open("chapters.json", "w") as chapters_file:
        json.dump(json.loads(response.choices[0].message.content), chapters_file)

def create_chapter_content(ebook_title, chapter, subheading):
    response = client.chat.completions.create(
        model=MODEL,
        max_tokens=3000,
        messages=[
            {"role": "system", "content": f"You are a creative eBook writer. The title of the eBook you are writing is '{ebook_title}'."},
            {"role": "user", "content": f"Write the text content for the subheading titled '{subheading}' under the chapter titled '{chapter}'. Be elaborate and clear. Formal language is preferred. Include the chapter name and subheading in the response. The content should be between 350 and 500 words long."} #Change the word limit as desired
        ]
    )
    return response.choices[0].message.content

def convert_docx_to_pdf(docx_path, pdf_path):
    if not os.path.exists(docx_path):
        print(f"The file {docx_path} does not exist.")
        return
    try:
        convert(docx_path, pdf_path)
        print(f"Conversion successful! PDF saved at: {pdf_path}")
    except Exception as e:
        print(f"An error occurred: {e}")

def create_audiobook_from_pdf(pdf_path, output_audio_path):
    reader = PdfReader(pdf_path)
    book = ""

    for page in reader.pages:
        text = page.extract_text()
        book += text

    def split_text(text, limit=3000):
        chunks = []
        while len(text) > limit:
            split_at = text.rfind(' ', 0, limit)
            if split_at == -1:
                split_at = limit
            chunks.append(text[:split_at])
            text = text[split_at:]
        chunks.append(text)
        return chunks

    chunks = split_text(book, limit=3000)

    output_dir = gettempdir()
    audio_files = []

    for i, chunk in enumerate(chunks):
        speech_file_path = Path(output_dir) / f"speech_{i}.mp3"
        with client.audio.speech.with_streaming_response.create(
            model="tts-1",
            voice="onyx",
            input=chunk,
        ) as response:
            response.stream_to_file(speech_file_path)
        audio_files.append(speech_file_path)


    combined_audio = AudioSegment.empty()
    for file in audio_files:
        combined_audio += AudioSegment.from_mp3(file)

    combined_audio.export(output_audio_path, format="mp3")
    
    # #Open the combined audio file
    # if sys.platform == "win32": #Windows
    #     os.startfile(output_audio_path)
    # elif sys.platform == "darwin": #macOS
    #     subprocess.call(["open", output_audio_path])
    # elif sys.platform == "linux": #Linux
    #     subprocess.call(["xdg-open", output_audio_path])

    print(f"Combined audio saved at {output_audio_path}")

def generate_ebook(title, base_dir):
    doc = Document()
    
    create_chapters_for_title(title)
    chapters = None
    with open("chapters.json") as chapters_file:
        chapters = json.load(chapters_file)
    
    for chapter, subheadings in chapters.items():
        doc.add_heading(chapter, level=1)
        for subheading in subheadings:
            contents = create_chapter_content(title, chapter, subheading)
            doc.add_heading(subheading, level=2)
            doc.add_paragraph(contents)

    sanitized_title = re.sub(r'[\\/*?:"<>|]', "_", title)
    docx_path = os.path.join(base_dir, f"{sanitized_title}.docx")
    pdf_path = os.path.join(base_dir, f"{sanitized_title}.pdf")
    audio_path = os.path.join(base_dir, f"{sanitized_title}.mp3")

    doc.save(docx_path)
    
    convert_docx_to_pdf(docx_path, pdf_path)
    create_audiobook_from_pdf(pdf_path, audio_path)
    
    if os.path.exists(docx_path):
        os.remove(docx_path)
        
    chapters_json_path = "chapters.json"
    if os.path.exists(chapters_json_path):
        os.remove(chapters_json_path)
    
    print(f"Completed generating the eBook '{title}'. PDF saved at: {pdf_path}. Audiobook saved at: {audio_path}.")
    
    
#Example Usage
# generate_ebook("World War 2 Timeline", os.path.dirname(os.path.abspath(__file__)))